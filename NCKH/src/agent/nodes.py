from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

import chromadb
from docx import Document
from llama_index.core import Document as LlamaDocument
from llama_index.core import Settings, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.llms.groq import Groq
from llama_index.vector_stores.chroma import ChromaVectorStore

from .state import AgentConfig


def extract_metadata_from_line(text_line: str) -> dict[str, str]:
    pattern = r"\|\s*(\d+)\.?\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|\s*(\d+)\s*\|\s*([^|]+)"
    match = re.search(pattern, text_line)
    if not match:
        return {}
    return {
        "ma_mon_hoc": match.group(2).strip(),
        "ten_mon_hoc": match.group(3).strip(),
        "so_tin_chi": match.group(4).strip(),
        "hoc_ky": match.group(5).strip().replace("\n", "").replace("|", "").strip(),
    }


def _iter_docx_files(data_dir: Path) -> Iterable[Path]:
    if not data_dir.exists():
        raise FileNotFoundError(f"Data directory does not exist: {data_dir}")
    files = sorted(p for p in data_dir.glob("*.docx") if not p.name.startswith("~$"))
    if not files:
        raise FileNotFoundError(f"No .docx files found in: {data_dir}")
    return files


def load_documents(data_dir: Path) -> list[LlamaDocument]:
    docs: list[LlamaDocument] = []
    for file_path in _iter_docx_files(data_dir):
        try:
            doc = Document(file_path)
        except Exception:
            continue
        file_name = file_path.name

        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) < 5:
                continue
            meta = {"file_source": file_name}
            meta.update(extract_metadata_from_line(text))
            docs.append(LlamaDocument(text=text, metadata=meta))

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells).strip()
                if len(row_text) < 5:
                    continue
                meta = {"file_source": file_name}
                meta.update(extract_metadata_from_line(row_text))
                docs.append(LlamaDocument(text=row_text, metadata=meta))
    return docs


def _configure_embedding(
    config: AgentConfig,
    forced_provider: str | None = None,
    forced_model: str | None = None,
) -> None:
    provider = forced_provider or config.embedding_provider
    model_name = forced_model or config.embedding_model

    if provider != "huggingface":
        raise ValueError(
            "Unsupported EMBEDDING_PROVIDER. Use 'huggingface' when running with Groq."
        )

    from llama_index.embeddings.huggingface import HuggingFaceEmbedding

    hf_model = model_name or "bkai-foundation-models/vietnamese-bi-encoder"
    Settings.embed_model = HuggingFaceEmbedding(model_name=hf_model)


def _get_embed_dimension() -> int | None:
    try:
        embed_model = Settings.embed_model
        if embed_model is None:
            return None
        vector = embed_model.get_text_embedding("test dimension")
        if not vector:
            return None
        return int(len(vector))
    except Exception:
        return None


def _configure_llm(config: AgentConfig) -> None:
    candidate_models = [
        config.groq_model,
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "llama-3.1-70b-versatile",
    ]
    last_error: Exception | None = None
    for model_name in dict.fromkeys(candidate_models):
        try:
            Settings.llm = Groq(
                model=model_name,
                api_key=config.groq_api_key,
                temperature=config.temperature,
            )
            return
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(
        "Cannot initialize Groq LLM with known model names. "
        "Set GROQ_MODEL explicitly and verify your API key."
    ) from last_error


def _get_collection_dimension(collection: object) -> int | None:
    try:
        result = collection.get(limit=1, include=["embeddings"])
        embeddings = result.get("embeddings")
        if embeddings is None or len(embeddings) == 0:
            return None
        return int(len(embeddings[0]))
    except Exception:
        return None


def get_or_build_index(config: AgentConfig, rebuild: bool = False) -> VectorStoreIndex:
    _configure_llm(config)
    _configure_embedding(config=config)
    config.vector_store_dir.mkdir(parents=True, exist_ok=True)
    db = chromadb.PersistentClient(path=str(config.vector_store_dir))
    if rebuild:
        try:
            db.delete_collection(config.collection_name)
        except Exception:
            pass
    collection = db.get_or_create_collection(config.collection_name)
    vector_store = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    has_data = collection.count() > 0
    existing_dim = _get_collection_dimension(collection) if has_data else None
    current_embed_dim = _get_embed_dimension() if has_data else None

    if has_data and not rebuild:
        if (
            existing_dim is not None
            and current_embed_dim is not None
            and existing_dim != current_embed_dim
        ):
            raise RuntimeError(
                f"Vector store dimension mismatch (store={existing_dim}, current={current_embed_dim}). "
                "Run with --rebuild to rebuild index with current embedding model."
            )

    if has_data and not rebuild:
        return VectorStoreIndex.from_vector_store(
            vector_store=vector_store,
            storage_context=storage_context,
        )

    documents = load_documents(config.data_dir)
    splitter = SentenceSplitter(
        chunk_size=config.chunk_size,
        chunk_overlap=config.chunk_overlap,
        paragraph_separator="\n",
        secondary_chunking_regex=r"[^.!?]+[.!?]?",
    )
    nodes = splitter.get_nodes_from_documents(documents)
    return VectorStoreIndex(nodes=nodes, storage_context=storage_context)
